"""DOCX document editor with diff preview and approval workflow."""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from .base import (
    ApprovalTracker,
    DocumentEditor,
    DocumentOperation,
    OperationResult,
    OperationType,
)
from .diff_display import render_text_diff, DIM, RESET


@dataclass
class DocxOperation(DocumentOperation):
    """DOCX-specific operation with additional context."""
    
    paragraph_index: int | None = None
    run_index: int | None = None
    search_pattern: str | None = None  # For find-replace operations


class DocxEditor(DocumentEditor):
    """Editor for DOCX files with tracked changes preview."""
    
    def get_supported_extensions(self) -> list[str]:
        return [".docx"]
    
    def render_diff(self, operation: DocumentOperation) -> str:
        """Render a human-readable diff for DOCX operations."""
        diff = render_text_diff(operation.old_text, operation.new_text)
        
        parts = []
        
        # Add location info
        if operation.location:
            parts.append(f"{DIM}@ {operation.location}{RESET}")
        
        # Add context before
        if operation.context_before:
            for line in operation.context_before.splitlines():
                parts.append(f"{DIM}  {line}{RESET}")
        
        # Add the diff
        parts.append(diff)
        
        # Add context after
        if operation.context_after:
            for line in operation.context_after.splitlines():
                parts.append(f"{DIM}  {line}{RESET}")
        
        return "\n".join(parts)
    
    def apply_operation(self, operation: DocumentOperation) -> OperationResult:
        """Apply a DOCX operation."""
        try:
            from docx import Document
        except ImportError:
            return OperationResult(
                success=False,
                output="python-docx not installed. Run: pip install python-docx",
            )
        
        path = self._resolve(operation.path)
        
        if operation.type == OperationType.DOCX_REPLACE_TEXT:
            return self._replace_text(path, operation)
        elif operation.type == OperationType.DOCX_INSERT_TEXT:
            return self._insert_text(path, operation)
        elif operation.type == OperationType.DOCX_DELETE_TEXT:
            return self._delete_text(path, operation)
        elif operation.type == OperationType.DOCX_ADD_COMMENT:
            return self._add_comment(path, operation)
        else:
            return OperationResult(
                success=False,
                output=f"Unsupported operation type: {operation.type}",
            )
    
    def _replace_text(self, path: Path, operation: DocumentOperation) -> OperationResult:
        """Replace text in a DOCX file."""
        from docx import Document
        
        if not path.exists():
            return OperationResult(success=False, output=f"File not found: {path}")
        
        doc = Document(str(path))
        old_text = operation.old_text or ""
        new_text = operation.new_text or ""
        
        replacements = 0
        
        # Search and replace in paragraphs
        for para in doc.paragraphs:
            if old_text in para.text:
                # Replace in runs to preserve formatting
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replacements += 1
                
                # If not found in individual runs, text might span runs
                # In that case, we need a more complex approach
                if replacements == 0 and old_text in para.text:
                    # Fallback: rebuild paragraph text
                    full_text = para.text
                    if old_text in full_text:
                        new_full = full_text.replace(old_text, new_text)
                        # Clear runs and add new text
                        for run in para.runs[1:]:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = new_full
                        replacements += 1
        
        # Also search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    replacements += 1
        
        if replacements == 0:
            return OperationResult(
                success=False,
                output=f"Text not found: '{old_text[:50]}...' " if len(old_text) > 50 else f"Text not found: '{old_text}'",
            )
        
        doc.save(str(path))
        return OperationResult(
            success=True,
            output=f"Replaced {replacements} occurrence(s) in {path.name}",
        )
    
    def _insert_text(self, path: Path, operation: DocumentOperation) -> OperationResult:
        """Insert text at a specific location."""
        from docx import Document
        
        if not path.exists():
            return OperationResult(success=False, output=f"File not found: {path}")
        
        doc = Document(str(path))
        new_text = operation.new_text or ""
        
        if isinstance(operation, DocxOperation) and operation.paragraph_index is not None:
            # Insert at specific paragraph
            idx = operation.paragraph_index
            if 0 <= idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                para.add_run(new_text)
            else:
                return OperationResult(
                    success=False,
                    output=f"Paragraph index {idx} out of range (0-{len(doc.paragraphs)-1})",
                )
        else:
            # Append to end
            doc.add_paragraph(new_text)
        
        doc.save(str(path))
        return OperationResult(
            success=True,
            output=f"Inserted text in {path.name}",
        )
    
    def _delete_text(self, path: Path, operation: DocumentOperation) -> OperationResult:
        """Delete text from document."""
        # Deletion is just replacement with empty string
        operation.new_text = ""
        return self._replace_text(path, operation)
    
    def _add_comment(self, path: Path, operation: DocumentOperation) -> OperationResult:
        """Add a comment to the document."""
        # Comments in python-docx are complex; for now, return info
        return OperationResult(
            success=False,
            output="Comment insertion not yet implemented. Use add_docx_comment tool directly.",
        )
    
    # Helper methods for creating operations
    
    @classmethod
    def create_replace_operation(
        cls,
        path: str,
        old_text: str,
        new_text: str,
        description: str = "",
        location: str | None = None,
        context_before: str | None = None,
        context_after: str | None = None,
    ) -> DocxOperation:
        """Create a text replacement operation."""
        return DocxOperation(
            type=OperationType.DOCX_REPLACE_TEXT,
            path=path,
            description=description or f"Replace text",
            old_text=old_text,
            new_text=new_text,
            location=location,
            context_before=context_before,
            context_after=context_after,
        )
    
    @classmethod
    def create_insert_operation(
        cls,
        path: str,
        new_text: str,
        paragraph_index: int | None = None,
        description: str = "",
        location: str | None = None,
    ) -> DocxOperation:
        """Create a text insertion operation."""
        return DocxOperation(
            type=OperationType.DOCX_INSERT_TEXT,
            path=path,
            description=description or "Insert text",
            new_text=new_text,
            location=location,
            paragraph_index=paragraph_index,
        )
    
    @classmethod
    def create_delete_operation(
        cls,
        path: str,
        old_text: str,
        description: str = "",
        location: str | None = None,
        context_before: str | None = None,
        context_after: str | None = None,
    ) -> DocxOperation:
        """Create a text deletion operation."""
        return DocxOperation(
            type=OperationType.DOCX_DELETE_TEXT,
            path=path,
            description=description or "Delete text",
            old_text=old_text,
            location=location,
            context_before=context_before,
            context_after=context_after,
        )
